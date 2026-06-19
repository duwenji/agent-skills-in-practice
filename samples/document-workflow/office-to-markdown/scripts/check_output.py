#!/usr/bin/env python3
"""
office-to-markdown 変換後の自動検証スクリプト。

変換結果の Markdown を検査し、構造欠落・画像プレースホルダー・識別子の
有無を報告する。.docx が元ファイルの場合は python-docx を使って
見出し数・テーブル数を原本と比較する。

使い方:
  python check_output.py <元ファイル> <変換後.md>
  python check_output.py 仕様書.docx 仕様書.md
"""

import re
import sys
import argparse
from dataclasses import dataclass
from enum import Enum
from pathlib import Path

# Windows コンソールで絵文字を出力できるよう UTF-8 に統一する
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")


# ---------------------------------------------------------------------------
# 結果型
# ---------------------------------------------------------------------------

class Status(Enum):
    PASS = "✅ PASS"
    WARN = "⚠️  WARN"
    FAIL = "❌ FAIL"


@dataclass
class CheckResult:
    name: str
    status: Status
    message: str


# ---------------------------------------------------------------------------
# Markdown の解析ユーティリティ
# ---------------------------------------------------------------------------

def count_md_headings(text: str) -> dict[int, int]:
    counts = {1: 0, 2: 0, 3: 0}
    for line in text.splitlines():
        stripped = line.lstrip()
        if stripped.startswith("### "):
            counts[3] += 1
        elif stripped.startswith("## "):
            counts[2] += 1
        elif stripped.startswith("# "):
            counts[1] += 1
    return counts


def count_md_tables(text: str) -> int:
    """区切り行（|---|）の数でテーブル数を数える。"""
    return len(re.findall(r"^\|[-|: ]+\|$", text, re.MULTILINE))


def count_empty_placeholders(text: str) -> int:
    """空の画像プレースホルダー ![anything]() または ![]() を検出する。"""
    return len(re.findall(r"!\[.*?\]\(\s*\)", text))


def count_embedded_images(text: str) -> int:
    """data:image/...;base64,... 形式の埋め込み画像を検出する。"""
    return len(re.findall(r"!\[.*?\]\(data:image/[^;]+;base64,", text))


def find_identifiers(text: str) -> list[str]:
    """IF-001, REQ-023 のような識別子パターンを抽出する。"""
    return sorted(set(re.findall(r"\b[A-Z]{1,5}-\d{3,}\b", text)))


# ---------------------------------------------------------------------------
# 各チェック関数
# ---------------------------------------------------------------------------

def check_word_count(original_path: Path, md_text: str) -> CheckResult:
    original_size = original_path.stat().st_size
    md_chars = len(md_text.replace(" ", "").replace("\n", ""))

    if original_size == 0:
        return CheckResult("文字数チェック", Status.WARN, "元ファイルのサイズが 0 です。")

    ratio = md_chars / original_size

    if ratio < 0.001:
        return CheckResult(
            "文字数チェック", Status.FAIL,
            f"出力が極端に短い可能性があります "
            f"(元: {original_size:,} bytes, 出力: {md_chars:,} 文字, 比率: {ratio:.3%})",
        )
    if ratio < 0.01:
        return CheckResult(
            "文字数チェック", Status.WARN,
            f"出力がやや短い可能性があります "
            f"(元: {original_size:,} bytes, 出力: {md_chars:,} 文字, 比率: {ratio:.3%})",
        )
    return CheckResult(
        "文字数チェック", Status.PASS,
        f"文字数は妥当です (元: {original_size:,} bytes, 出力: {md_chars:,} 文字)",
    )


def check_headings(md_text: str) -> CheckResult:
    counts = count_md_headings(md_text)
    total = sum(counts.values())
    if total == 0:
        return CheckResult(
            "見出し構造チェック", Status.WARN,
            "見出し (#, ##, ###) が1件も検出されませんでした。"
            "元ファイルに見出しがない場合は正常です。",
        )
    return CheckResult(
        "見出し構造チェック", Status.PASS,
        f"見出しを検出: H1={counts[1]} 件, H2={counts[2]} 件, H3={counts[3]} 件",
    )


def check_tables(md_text: str) -> CheckResult:
    count = count_md_tables(md_text)
    if count == 0:
        return CheckResult(
            "テーブルチェック", Status.WARN,
            "テーブルが検出されませんでした。元ファイルに表がない場合は正常です。",
        )
    return CheckResult(
        "テーブルチェック", Status.PASS,
        f"テーブルを {count} 件検出しました。",
    )


def check_image_placeholders(md_text: str) -> CheckResult:
    count = count_empty_placeholders(md_text)
    if count > 0:
        return CheckResult(
            "画像プレースホルダーチェック", Status.WARN,
            f"空の画像プレースホルダーが {count} 件あります。"
            "グラフ・画像が変換されていません。処理方針 B または C への変更を検討してください。",
        )
    return CheckResult(
        "画像プレースホルダーチェック", Status.PASS,
        "空の画像プレースホルダーはありません。",
    )


def check_embedded_images(md_text: str) -> CheckResult:
    count = count_embedded_images(md_text)
    if count > 0:
        return CheckResult(
            "埋め込み画像チェック", Status.WARN,
            f"base64 埋め込み画像が {count} 件あります。"
            " --extract-images オプションを付けて再変換するとファイルとして抽出されます。",
        )
    return CheckResult(
        "埋め込み画像チェック", Status.PASS,
        "base64 埋め込み画像はありません。",
    )


def check_identifiers(md_text: str) -> CheckResult:
    ids = find_identifiers(md_text)
    if ids:
        preview = ", ".join(ids[:10]) + ("..." if len(ids) > 10 else "")
        return CheckResult(
            "識別子チェック", Status.PASS,
            f"識別子を {len(ids)} 種類検出: {preview}",
        )
    return CheckResult(
        "識別子チェック", Status.WARN,
        "識別子パターン (IF-001 等) が検出されませんでした。"
        "IF 仕様書の場合は欠落がないか確認してください。",
    )


def check_docx_structure(original_path: Path, md_text: str) -> list[CheckResult]:
    """python-docx を使って .docx の構造と変換結果を比較する。"""
    results: list[CheckResult] = []
    try:
        from docx import Document  # type: ignore
    except ImportError:
        results.append(CheckResult(
            "構造比較 (docx)", Status.WARN,
            "python-docx が未インストールのため詳細比較をスキップしました。"
            "`pip install python-docx` で有効化できます。",
        ))
        return results

    try:
        doc = Document(original_path)
    except Exception as e:
        results.append(CheckResult(
            "構造比較 (docx)", Status.WARN,
            f"元ファイルの読み込みに失敗しました: {e}",
        ))
        return results

    # 見出し数の比較
    src_headings = sum(1 for p in doc.paragraphs if p.style.name.startswith("Heading"))
    md_headings = sum(count_md_headings(md_text).values())

    if src_headings > 0:
        ratio = md_headings / src_headings
        if ratio < 0.7:
            results.append(CheckResult(
                "見出し数比較 (docx)", Status.FAIL,
                f"見出し数が大幅に減少しています "
                f"(元: {src_headings} 件 → 変換後: {md_headings} 件, 残存率: {ratio:.0%})",
            ))
        else:
            results.append(CheckResult(
                "見出し数比較 (docx)", Status.PASS,
                f"見出し数: 元ファイル {src_headings} 件 → 変換後 {md_headings} 件",
            ))

    # テーブル数の比較
    src_tables = len(doc.tables)
    md_tables = count_md_tables(md_text)

    if src_tables > 0:
        if md_tables < src_tables:
            results.append(CheckResult(
                "テーブル数比較 (docx)", Status.WARN,
                f"テーブル数が減少しています "
                f"(元: {src_tables} 件 → 変換後: {md_tables} 件)",
            ))
        else:
            results.append(CheckResult(
                "テーブル数比較 (docx)", Status.PASS,
                f"テーブル数: 元ファイル {src_tables} 件 → 変換後 {md_tables} 件",
            ))

    return results


# ---------------------------------------------------------------------------
# レポート出力
# ---------------------------------------------------------------------------

def run_checks(original_path: Path, md_path: Path) -> list[CheckResult]:
    md_text = md_path.read_text(encoding="utf-8")
    results: list[CheckResult] = []

    results.append(check_word_count(original_path, md_text))
    results.append(check_headings(md_text))
    results.append(check_tables(md_text))
    results.append(check_image_placeholders(md_text))
    results.append(check_embedded_images(md_text))
    results.append(check_identifiers(md_text))

    if original_path.suffix.lower() == ".docx":
        results.extend(check_docx_structure(original_path, md_text))

    return results


def print_report(results: list[CheckResult]) -> int:
    """レポートを出力し、FAIL があれば exit code 1 を返す。"""
    print("\n" + "=" * 60)
    print("変換後チェックレポート")
    print("=" * 60)

    fail_count = 0
    warn_count = 0

    for r in results:
        print(f"\n{r.status.value}  {r.name}")
        print(f"   {r.message}")
        if r.status == Status.FAIL:
            fail_count += 1
        elif r.status == Status.WARN:
            warn_count += 1

    print("\n" + "-" * 60)
    print(f"結果: FAIL={fail_count}件  WARN={warn_count}件  PASS={len(results) - fail_count - warn_count}件")

    if fail_count > 0:
        print("❌ 変換内容を確認してください。受領原本を正本として Markdown を修正してください。")
        return 1
    if warn_count > 0:
        print("⚠️  警告を確認し、問題なければそのまま進めてください。")
    else:
        print("✅ すべてのチェックをパスしました。")
    return 0


# ---------------------------------------------------------------------------
# エントリーポイント
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Office → Markdown 変換後の自動検証",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument("original", help="変換元の Office ファイルパス")
    parser.add_argument("markdown", help="変換後の Markdown ファイルパス")
    args = parser.parse_args()

    original_path = Path(args.original)
    md_path = Path(args.markdown)

    for p in (original_path, md_path):
        if not p.exists():
            print(f"エラー: ファイルが見つかりません: {p}", file=sys.stderr)
            sys.exit(1)

    results = run_checks(original_path, md_path)
    exit_code = print_report(results)
    sys.exit(exit_code)


if __name__ == "__main__":
    main()
